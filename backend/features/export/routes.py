"""
Document export endpoints for Word and PDF generation.
"""

from flask import Blueprint, request, send_file, jsonify
from io import BytesIO
from datetime import datetime

export_bp = Blueprint("export", __name__)


@export_bp.route("/docx", methods=["POST"])
def export_docx():
    """
    Export document as Word (.docx) file.

    Request body:
        - content: Document content (markdown/text)
        - title: Document title (optional)
        - metadata: Document metadata (optional)

    Returns:
        Word document file download
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500

    data = request.get_json()
    content = data.get("content", "")
    title = data.get("title", "Legal Document")
    metadata = data.get("metadata", {})

    if not content:
        return jsonify({"error": "No content provided"}), 400

    # Create Word document
    doc = Document()

    # Add title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add metadata if provided
    if metadata:
        doc.add_paragraph()
        meta_text = []
        if metadata.get("party_a"):
            meta_text.append(f"Party A: {metadata['party_a']}")
        if metadata.get("party_b"):
            meta_text.append(f"Party B: {metadata['party_b']}")
        if metadata.get("effective_date"):
            meta_text.append(f"Effective Date: {metadata['effective_date']}")
        if meta_text:
            for text in meta_text:
                p = doc.add_paragraph(text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Parse and add content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        else:
            doc.add_paragraph(line)

    # Add footer with generation info
    doc.add_paragraph()
    footer = doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(8)
    footer_run.italic = True

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"

    return send_file(
        buffer,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )


@export_bp.route("/pdf", methods=["POST"])
def export_pdf():
    """
    Export document as PDF file.

    Request body:
        - content: Document content (markdown/text)
        - title: Document title (optional)
        - metadata: Document metadata (optional)

    Returns:
        PDF document file download
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    except ImportError:
        return jsonify({"error": "reportlab not installed"}), 500

    data = request.get_json()
    content = data.get("content", "")
    title = data.get("title", "Legal Document")
    metadata = data.get("metadata", {})

    if not content:
        return jsonify({"error": "No content provided"}), 400

    # Create PDF in memory
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )

    # Styles
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        alignment=TA_CENTER,
        spaceAfter=20
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10
    )

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_JUSTIFY,
        spaceBefore=6,
        spaceAfter=6
    )

    meta_style = ParagraphStyle(
        'MetaStyle',
        parent=styles['Normal'],
        fontSize=10,
        alignment=TA_CENTER,
        spaceAfter=3
    )

    footer_style = ParagraphStyle(
        'FooterStyle',
        parent=styles['Normal'],
        fontSize=8,
        alignment=TA_CENTER,
        textColor='gray'
    )

    # Build content
    story = []

    # Title
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2 * inch))

    # Metadata
    if metadata:
        if metadata.get("party_a"):
            story.append(Paragraph(f"Party A: {metadata['party_a']}", meta_style))
        if metadata.get("party_b"):
            story.append(Paragraph(f"Party B: {metadata['party_b']}", meta_style))
        if metadata.get("effective_date"):
            story.append(Paragraph(f"Effective Date: {metadata['effective_date']}", meta_style))
        story.append(Spacer(1, 0.3 * inch))

    # Parse content
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.1 * inch))
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], title_style))
        elif line.startswith('## ') or line.startswith('### '):
            clean_line = line.lstrip('#').strip()
            story.append(Paragraph(clean_line, heading_style))
        elif line.startswith('**') and line.endswith('**'):
            story.append(Paragraph(f"<b>{line[2:-2]}</b>", body_style))
        else:
            # Escape special characters for reportlab
            safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_line, body_style))

    # Footer
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(
        f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        footer_style
    ))

    # Build PDF
    doc.build(story)
    buffer.seek(0)

    filename = f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"

    return send_file(
        buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=filename
    )


@export_bp.route("/preview", methods=["POST"])
def preview_document():
    """
    Return formatted document content for preview.

    Request body:
        - content: Document content
        - format: Output format (html, markdown, plain)

    Returns:
        Formatted document content
    """
    data = request.get_json()
    content = data.get("content", "")
    output_format = data.get("format", "html")

    if not content:
        return jsonify({"error": "No content provided"}), 400

    if output_format == "html":
        # Simple markdown to HTML conversion
        html = content
        html = html.replace('\n\n', '</p><p>')
        html = html.replace('\n', '<br>')

        # Headers
        import re
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)

        # Bold
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)

        html = f"<div class='document-preview'><p>{html}</p></div>"

        return jsonify({"content": html, "format": "html"})

    elif output_format == "markdown":
        return jsonify({"content": content, "format": "markdown"})

    else:  # plain
        # Remove markdown formatting
        plain = content
        plain = plain.replace('**', '')
        plain = plain.replace('# ', '')
        plain = plain.replace('## ', '')
        plain = plain.replace('### ', '')

        return jsonify({"content": plain, "format": "plain"})
